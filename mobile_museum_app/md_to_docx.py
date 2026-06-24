# -*- coding: utf-8 -*-
"""将 Markdown 文件转换为 Word (.docx) 文档。"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table


def set_cn_font(run, name: str = "宋体", size: int | None = None, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def add_runs_with_inline(paragraph, text: str, base_size: int = 11, bold_base: bool = False):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_cn_font(run, size=base_size, bold=bold_base)
        chunk = m.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_cn_font(run, size=base_size, bold=True)
        elif chunk.startswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            set_cn_font(run, size=base_size, bold=bold_base)
            run.italic = True
        elif chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            set_cn_font(run, name="Consolas", size=base_size)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_cn_font(run, size=base_size, bold=bold_base)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_with_inline(p, cell_text, base_size=10, bold_base=(i == 0))
    doc.add_paragraph("")


def add_code_block(doc: Document, lines: list[str]):
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_cn_font(run, name="Consolas", size=9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph("")


def md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # 默认正文样式
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []
    list_buffer: list[tuple[int, str]] = []

    def flush_list():
        nonlocal list_buffer
        for level, item in list_buffer:
            p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            add_runs_with_inline(p, item, base_size=11)
        list_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_list()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            flush_list()
            i += 1
            continue

        if stripped == "---":
            flush_list()
            p = doc.add_paragraph("—" * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("#"):
            flush_list()
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            heading = doc.add_heading("", level=min(level, 4))
            add_runs_with_inline(heading, title, base_size=16 - level, bold_base=True)
            i += 1
            continue

        if stripped.startswith(">"):
            flush_list()
            quote = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            add_runs_with_inline(p, quote, base_size=11)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            flush_list()
            table_rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        m = re.match(r"^(\s*)-\s+(.+)$", line)
        if m:
            indent = len(m.group(1))
            level = 1 if indent >= 2 else 0
            list_buffer.append((level, m.group(2)))
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            flush_list()
            p = doc.add_paragraph(style="List Number")
            add_runs_with_inline(p, m.group(2), base_size=11)
            i += 1
            continue

        flush_list()
        p = doc.add_paragraph()
        add_runs_with_inline(p, stripped, base_size=11)
        i += 1

    flush_list()
    if in_code and code_lines:
        add_code_block(doc, code_lines)

    doc.save(str(docx_path))
    print(f"OK: {docx_path}")


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("用法: python md_to_docx.py file1.md [file2.md ...]")
        return 1
    for arg in argv[1:]:
        md = Path(arg)
        if not md.exists():
            print(f"跳过（不存在）: {md}")
            continue
        out = md.with_suffix(".docx")
        md_to_docx(md, out)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
